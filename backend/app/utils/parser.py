"""
文档解析器
支持 .docx、.pdf、.txt、.md 格式
"""
import os
from typing import List, Dict
import pdfplumber
import docx
from app.models.document import Document

class DocumentParser:
    """文档解析器类，支持 .docx、.pdf、.txt、.md 格式"""
    
    def parse(self, file_path: str) -> Document:
        """
        解析文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            Document 对象
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        file_name = os.path.basename(file_path)
        file_type = self._get_file_type(file_path)
        
        pages = []
        if file_type == "docx":
            content, sections = self._parse_docx(file_path)
        elif file_type == "pdf":
            content, sections, pages = self._parse_pdf_with_pages(file_path)
        elif file_type in ["txt", "md"]:
            content, sections = self._parse_text(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
            
        return Document(
            document_name=file_name,
            file_path=file_path,
            file_type=file_type,
            content=content,
            pages=pages,
            sections=sections
        )

    def _parse_pdf_with_pages(self, file_path: str) -> (str, List[str], List[Dict]):
        """解析 PDF 并保留页码信息"""
        content = []
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    content.append(text)
                    pages.append({"page_num": i + 1, "text": text})
        return "\n".join(content), [], pages
    
    def _get_file_type(self, file_path: str) -> str:
        """获取文件类型"""
        ext = file_path.split(".")[-1].lower()
        return ext
    
    def _parse_docx(self, file_path: str) -> (str, List[str]):
        """
        解析 .docx 文件
        使用 python-docx 提取文本
        """
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text), []
    
    def _parse_text(self, file_path: str) -> (str, List[str]):
        """
        解析 .txt 或 .md 文件
        直接读取文本
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content, []

